from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\AMTL\repos\ck-life-os")
EVIDENCE_DIR = Path(
    r"C:\Users\Mani\AMTL-docs\AMTL-AGENT-CONTROL-CENTER\evidence\ck-life-os-local-internal-270626"
)
DOC_VERSION = "v2.0"
DOC_DATE = "28 June 2026"
OUT = ROOT / f"CK-Life-OS-Complete-Story-Flow-Guide-{DOC_VERSION}-20260628.docx"

BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
GOLD = "C9944A"
PALE_GOLD = "FFF4DD"
GREEN = "E8F4EF"
INK = RGBColor(20, 31, 43)
MUTED = RGBColor(91, 103, 116)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = "D7DBE2") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margin = tc_pr.first_child_found_in("w:tcMar")
    if margin is None:
        margin = OxmlElement("w:tcMar")
        tc_pr.append(margin)
    for m, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margin.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            margin.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[float]) -> None:
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
            set_cell_margins(row.cells[idx])
            set_cell_border(row.cells[idx])
            row.cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_run(run, *, size=None, bold=None, color=None, italic=None) -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph(doc, text="", *, style=None, bold=False, color=None, size=None, italic=False, align=None, before=0, after=6):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.18
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run(r, size=size, bold=bold, color=color, italic=italic)
    return p


def heading(doc, text: str, level: int = 1):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.bold = True
    r.font.name = "Calibri"
    r._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    r._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    r.font.color.rgb = RGBColor.from_string(BLUE if level < 3 else "1F3A5F")
    r.font.size = Pt({1: 16, 2: 13, 3: 12}.get(level, 11))
    p.paragraph_format.space_before = Pt({1: 18, 2: 14, 3: 10}.get(level, 8))
    p.paragraph_format.space_after = Pt({1: 10, 2: 7, 3: 5}.get(level, 4))
    return p


def callout(doc, title: str, body: str, fill: str = PALE_GOLD, spacer: bool = True):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run(r, bold=True, color=BLUE, size=11)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(body)
    set_run(r2, size=10.5, color="202A36")
    if spacer:
        paragraph(doc, "", after=3)


def table(doc, headers: list[str], rows: list[list[str]], widths: list[float], header_fill: str = LIGHT_BLUE):
    t = doc.add_table(rows=1, cols=len(headers))
    set_table_width(t, widths)
    for i, h in enumerate(headers):
        cell = t.cell(0, i)
        set_cell_shading(cell, header_fill)
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run(r, bold=True, color=BLUE, size=9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run(r, size=9.2, color="202A36")
            set_cell_margins(cells[i])
            set_cell_border(cells[i])
    paragraph(doc, "", after=3)
    return t


def flow(doc, title: str, steps: list[str]):
    callout(doc, title, " -> ".join(steps), fill=GREEN)


def bullet(doc, text: str):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    set_run(r, size=10.5, color="202A36")


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(0.85))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color in (("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, "1F3A5F")):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    return doc


def add_cover(doc: Document) -> None:
    paragraph(doc, "Almost Magic Tech Lab", bold=True, color=GOLD, size=11, after=4)
    paragraph(doc, "CK / Life OS", bold=True, color=BLUE, size=26, after=2)
    paragraph(doc, f"Complete Story-Flow Guide {DOC_VERSION}", color="3F4E5F", size=16, after=14)
    paragraph(
        doc,
        "A story-like operating map of every visible module, screen, tab, panel, button, field, backend route, data boundary, and proof path in the local/internal CK / Life OS app.",
        size=11.5,
        after=14,
    )
    table(
        doc,
        ["Field", "Value"],
        [
            ["Product", "CK / Life OS"],
            ["Front door", "http://amtl/ck-life-os/ and http://127.0.0.1:5420/"],
            ["Runtime scope", "Local/internal app, port 5420"],
            ["Primary design rule", "Calm, tabbed, low-pressure, neurodivergent-friendly flow"],
            ["Version", DOC_VERSION],
            ["Document generated", DOC_DATE],
            ["Regeneration reason", "Updated after RAG / Sources source-draft changes, approved external URL fetch boundary, Product Bible updates, UI changes, and OpenRouter free/modest/expensive reasoning ladder."],
        ],
        [1.55, 4.95],
        header_fill=PALE_GOLD,
    )
    screenshot = EVIDENCE_DIR / "ck-life-os-calm-tabs-desktop.png"
    if screenshot.exists():
        paragraph(doc, "Current calm tabbed layout screenshot", bold=True, color=BLUE, after=4)
        doc.add_picture(str(screenshot), width=Inches(6.35))


def add_global_orientation(doc: Document) -> None:
    heading(doc, "1. Opening The App", 1)
    paragraph(
        doc,
        "The current local/internal product does not present a public login form. The user opens the front door directly. The story begins when the user visits http://amtl/ck-life-os/ or http://127.0.0.1:5420/.",
    )
    flow(
        doc,
        "Opening flow",
        [
            "User opens CK / Life OS",
            "Top bar confirms AMTL, CK Life OS, date/time, OpenRouter status, and local seal",
            "Left grouped menu appears",
            "Middle workspace shows the selected tabbed screen",
            "Right context rail shows only the relevant next action by default",
        ],
    )
    table(
        doc,
        ["Area", "What the user sees", "Why it matters"],
        [
            ["Top bar", "AMTL logo, CK Life OS logo/name, today's signal, date/time, OpenRouter readiness, AMTL local/internal seal", "Establishes identity, local runtime truth, model-provider readiness, and calm orientation."],
            ["Left menu", "Practice, Private, Knowledge, Reports, Admin / Proof", "Groups actions by user intent and keeps inactive areas collapsed."],
            ["Middle workspace", "Tabbed screens with one visible content panel at a time", "Prevents the user from scrolling through long stacked columns."],
            ["Right context rail", "Next action open; Selected detail, 6W / ELI10, Receipt / Evidence collapsed", "Shows relevant help without flooding the screen."],
        ],
        [1.25, 3.0, 2.25],
    )

    heading(doc, "2. Left Menu Categories", 1)
    table(
        doc,
        ["Menu category", "Menu items", "Story purpose"],
        [
            ["Practice", "Practice desk; Guide me; Innovation lens", "Daily practice and low-pressure guidance."],
            ["Private", "Encrypted journal; Inner work", "Private reflection, encrypted storage, and guided self-inquiry."],
            ["Knowledge", "RAG / Sources", "Local RAG/source status, source intake, approved NAS search, and pasted source-text search."],
            ["Reports", "Reports and gates", "Local summary, markdown export, staged gate, and synthetic-data checks."],
            ["Admin / Proof", "Evidence summaries; Delivery matrix; R2D2 equivalent; Dependency fallback; OpenRouter routing; OSS adoption", "Proof surfaces for review, kept out of normal practice flow."],
        ],
        [1.35, 2.65, 2.5],
    )
    callout(
        doc,
        "Version alignment",
        "This v2.0 guide uses the current visible UI names from index.html and the current route names from main.py. It includes the newer RAG / Sources draft/fetch behaviour and the OpenRouter free -> modest -> expensive reasoning ladder.",
        fill=GREEN,
    )


def add_practice_story(doc: Document) -> None:
    heading(doc, "3. Practice Desk Story", 1)
    paragraph(
        doc,
        "The user opens Practice -> Practice desk. The app shows the main story panel, a practice note field, two immediate buttons, today's truth cards, and practice tabs. The user can do one small practice without being scored.",
    )
    table(
        doc,
        ["Screen element", "Exact name", "What happens"],
        [
            ["Main heading", "Five practices for living", "Frames the product as a calm practice OS."],
            ["Field", "Practice note", "Optional note; blank is allowed."],
            ["Button", "Make today gentler", "Activates Difficult Month Mode locally."],
            ["Button", "Explain first step", "Opens Guide me for the practice screen."],
            ["Count card", "Practices", "Opens Guide me for practice orientation."],
            ["Count card", "Scores", "Explains intentional zero."],
            ["Count card", "Streaks", "Explains intentional zero."],
            ["Count card", "Lens items", "Opens Innovation lens and loads rows."],
        ],
        [1.3, 1.9, 3.3],
    )
    flow(
        doc,
        "Record a practice",
        [
            "User opens Practice desk",
            "User selects Presence / Reflection / Intention / Gratitude / Equanimity tab",
            "User writes optional Practice note",
            "User clicks Record Presence",
            "App calls POST /api/practices/presence/record",
            "Local JSONL receipt is written",
            "Right rail updates with next action, selected detail, 6W / ELI10, and receipt proof",
        ],
    )
    table(
        doc,
        ["Practice tab", "Buttons on card", "Backend routes", "Result"],
        [
            ["Presence", "Record Presence; Explain prompt; Details", "/api/practices/presence/record; /api/prompts/after-save/presence; /api/contextual-guide/presence", "Records awareness/presence receipt or explains the practice."],
            ["Reflection", "Record Reflection; Explain prompt; Details", "/api/practices/reflection/record; /api/prompts/after-save/reflection; /api/contextual-guide/reflection", "Records contemplation/reflection receipt or explains the practice."],
            ["Intention", "Record Intention; Explain prompt; Details", "/api/practices/intention/record; /api/prompts/after-save/intention; /api/contextual-guide/intention", "Records direction/values receipt or explains the practice."],
            ["Gratitude", "Record Gratitude; Explain prompt; Details", "/api/practices/gratitude/record; /api/prompts/after-save/gratitude; /api/contextual-guide/gratitude", "Records appreciation receipt or explains the practice."],
            ["Equanimity", "Record Equanimity; Explain prompt; Details", "/api/practices/equanimity/record; /api/prompts/after-save/equanimity; /api/contextual-guide/equanimity", "Records steadiness/balance receipt or explains the practice."],
        ],
        [1.1, 1.65, 2.25, 1.5],
    )
    callout(
        doc,
        "What the user learns",
        "CK / Life OS is deliberately not a scorekeeper. The zero score and zero streak cards are explanatory proof, not missing data.",
        fill=GREEN,
    )


def add_guide_and_lens(doc: Document) -> None:
    heading(doc, "4. Guide Me Story", 1)
    paragraph(doc, "The user opens Practice -> Guide me. The screen has two tabs: Start and Practice 6W.")
    table(
        doc,
        ["Tab or button", "Exact name", "What the app does"],
        [
            ["Tab", "Start", "Shows screen-specific guide choices."],
            ["Tab", "Practice 6W", "Shows each practice with contextual What / Who / Why / When / How / Where / ELI10."],
            ["Button", "Practice", "Calls /api/guided-use?view=practice."],
            ["Button", "Ideas", "Calls /api/guided-use?view=ideas."],
            ["Button", "Reports", "Calls /api/guided-use?view=reports."],
            ["Button", "Evidence", "Calls /api/guided-use?view=evidence."],
        ],
        [1.1, 1.6, 3.8],
    )
    flow(
        doc,
        "Tired-user guide flow",
        [
            "User opens Guide me",
            "User clicks Practice / Ideas / Reports / Evidence",
            "App returns start_here, next_click, evidence_to_check, do_not_do_yet, write_boundary",
            "Right rail updates with the selected guide",
        ],
    )

    heading(doc, "5. Innovation Lens Story", 1)
    paragraph(
        doc,
        "The user opens Practice -> Innovation lens. The lens turns the 500 generated field/navigation innovation patterns into calm, contextual helpers instead of a raw register.",
    )
    table(
        doc,
        ["Tab", "Fields or buttons", "Story"],
        [
            ["Filter", "Practice; Situation; Open matching innovations", "The user narrows the 500 patterns to a calm subset."],
            ["Results", "Open {idea_id}", "The app shows matching rows and lets the user open one pattern."],
            ["Detail", "Usefulness, control, cost, calm design, neurodivergent support, 6W / ELI10", "The user sees why the pattern matters and whether to use it."],
        ],
        [1.1, 2.5, 2.9],
    )
    flow(
        doc,
        "Innovation lens flow",
        [
            "User chooses Practice and Situation",
            "User clicks Open matching innovations",
            "App calls /api/ideas",
            "Results tab opens",
            "User clicks Open {idea_id}",
            "App calls /api/ideas/{idea_id}",
            "Detail tab opens with 6W / ELI10 and synthetic boundary",
        ],
    )
    table(
        doc,
        ["Situation filter", "Meaning in the story"],
        [
            ["morning", "Start the day gently."],
            ["midday", "Recover attention or reduce friction during the day."],
            ["evening", "Reflect without turning the day into performance."],
            ["difficult-day", "Use gentler supports when pressure is high."],
            ["low-energy", "Reduce effort and choice load."],
            ["family", "Make practice useful around relationships/family."],
            ["work", "Use the practice inside work context."],
            ["body", "Bring attention back to body and energy."],
            ["creative", "Support creativity without clutter."],
            ["repair", "Help after a rupture or difficult moment."],
        ],
        [1.5, 5.0],
    )


def add_private_story(doc: Document) -> None:
    heading(doc, "6. Encrypted Journal Story", 1)
    paragraph(doc, "The user opens Private -> Encrypted journal. The screen has tabs: Write, Entries, and Detail.")
    table(
        doc,
        ["Tab", "Fields/buttons", "What happens"],
        [
            ["Write", "Practice; Title; Encrypted journal entry; Save encrypted entry; Refresh journal", "User writes a private note. Save encrypts title/content with Fernet in the local runtime folder."],
            ["Entries", "Open local decrypt", "User sees encrypted entries and can open one locally."],
            ["Detail", "Decrypted content", "The app decrypts only for the local UI and shows the note."],
        ],
        [1.1, 2.7, 2.7],
    )
    flow(
        doc,
        "Journal save and read flow",
        [
            "User writes in Encrypted journal entry",
            "User clicks Save encrypted entry",
            "App calls POST /api/journal",
            "Encrypted JSONL row is written",
            "User clicks Refresh journal",
            "Entries tab lists local encrypted entries",
            "User clicks Open local decrypt",
            "Detail tab shows readable content inside local UI only",
        ],
    )

    heading(doc, "7. Inner Work Story", 1)
    paragraph(
        doc,
        "The user opens Private -> Inner work. This is for shadow integration, observation / Krishnamurti-style inquiry, grounding, relationship patterns, and unknown-mode reflection. It guides looking; it does not diagnose, therapise, or claim spiritual authority.",
    )
    table(
        doc,
        ["Tab", "Fields/buttons", "What happens"],
        [
            ["Start", "Mode; Intention; Inner work text input; Voice transcript input; Start guided session; Load sessions", "User chooses depth and writes text or pasted voice transcript. The app returns a posture, first question, and seven-step program."],
            ["Sessions", "Open session", "User loads encrypted session history and opens a selected session."],
            ["Detail", "Guidance detail", "Shows mode label, posture, not-answer boundary, first question, safety state, program, controls, AI provider called false."],
        ],
        [1.1, 2.9, 2.5],
    )
    table(
        doc,
        ["Mode", "Story use"],
        [
            ["Shadow integration", "Meet a rejected/protective part without acting it out or shaming it."],
            ["Observation / Krishnamurti-style inquiry", "Observe thought, fear, image, comparison, and becoming without authority claims."],
            ["Grounding first", "Stabilise before depth; choose safety over insight."],
            ["Relationship pattern", "Separate event, story, need, and next clean action."],
            ["I do not know yet", "Start gently and choose the right depth after the first reflection."],
        ],
        [2.0, 4.5],
    )
    flow(
        doc,
        "Inner Work safety flow",
        [
            "User writes text",
            "App checks for risk language",
            "If risk language appears, guidance switches to Grounding first",
            "Depth becomes grounding_only",
            "Session is encrypted locally",
            "External send and provider call remain false",
        ],
    )


def add_sources_reports_evidence(doc: Document) -> None:
    heading(doc, "8. RAG / Sources Story", 1)
    paragraph(doc, "The user opens Knowledge -> RAG / Sources. This is where RAG lives: approved internal/NAS source search, local source drafts, pasted source text, and explicitly approved one-time external URL fetches into local RAG. The app does not fetch a URL without the approval control.")
    table(
        doc,
        ["Tab", "Fields/buttons", "What happens"],
        [
            ["RAG status", "Explain RAG access; Refresh source status; Build local index", "Shows RAG location, internal/external lanes, approved roots, exact two exclusions, provider-call false, and pgvector-not-required truth."],
            ["Add source", "Source type; Title; Path or URL; RAG source notes; RAG source text; Approve one-time external URL fetch into local RAG; Stage source for approval", "Stages an internal path or external reference locally. Optional pasted source text becomes searchable local RAG material. External URLs fetch into local RAG only when the one-time approval checkbox is selected."],
            ["Search", "Search sources; Search", "User searches combined local RAG drafts and approved NAS source index."],
            ["Results", "Source matches", "Shows draft or NAS-index matches, local paths/references, and safe snippets."],
        ],
        [1.1, 2.5, 2.9],
    )
    callout(
        doc,
        "Approved NAS exclusions",
        "\\\\NAS2\\amtl-documents\\escalation-logs\\linux logs\\wazuz\\prefe\n\\\\Nas1\\Nas\\Programs\\Spiritual\\Dr Joe Dispenza\\Generous moment\\mani",
        fill=PALE_GOLD,
    )
    flow(
        doc,
        "RAG / Sources flow",
        [
            "User opens RAG / Sources",
            "User clicks RAG status",
            "App calls /api/rag/status",
            "User may stage internal/external source reference",
            "App calls POST /api/rag/source-draft",
            "Pasted source text is saved locally for search",
            "User may build local NAS index",
            "App calls POST /api/source-index/build",
            "User searches",
            "App calls /api/rag/search",
            "Results tab shows local draft and NAS matches",
        ],
    )

    heading(doc, "9. Reports And Gates Story", 1)
    paragraph(doc, "The user opens Reports -> Reports and gates. The section has Summary and Synthetic data tabs.")
    table(
        doc,
        ["Tab", "Button", "Backend route", "User result"],
        [
            ["Summary", "Open local report", "/api/reports/local-summary", "Shows report ID, receipt count, encrypted journal count, innovation count, external-send false, source-write false."],
            ["Summary", "Export markdown summary", "/api/reports/local-summary.md", "Shows copy-ready markdown in the local UI."],
            ["Summary", "Stage Elaine local gate", "/api/handoffs/lifeos-elaine-local/stage", "Creates a local staged receipt; nothing is sent to Elaine."],
            ["Synthetic data", "Check boundary", "/api/synthetic-data/status", "Shows synthetic row status and boundary truth."],
            ["Synthetic data", "Preview cleanup", "/api/synthetic-data/removal-preview", "Shows what would be removed; currently zero operational user history."],
            ["Synthetic data", "Run cleanup verification", "/api/synthetic-data/cleanup", "Creates local verification receipt proving synthetic boundary."],
        ],
        [1.1, 1.55, 1.85, 2.0],
    )

    heading(doc, "10. Evidence Summaries Story", 1)
    paragraph(doc, "The user opens Admin / Proof -> Evidence summaries. Proof is not shown in normal practice flow; it lives behind Checks and Output tabs.")
    table(
        doc,
        ["Button", "Route", "What it proves"],
        [
            ["Product Bible matrix", "/api/product-bible-matrix", "Feature coverage, delivery matrix, source index, and scope truth."],
            ["R2D2 equivalent", "/api/r2d2", "Requirement-to-delivery audit checks."],
            ["Dependency fallback", "/api/dependency-status", "Core UI does not require Postgres, scheduler, provider, source index, or unrelated apps."],
            ["Data truth", "/api/data-truth", "Receipts, encrypted journal, inner work, innovation catalogue, NAS source index, external-send/source-write/provider-call truth."],
            ["Button/drilldown truth", "/api/ui-truth", "Buttons, drilldowns, counts, and UI truth."],
            ["OSS adoption", "/api/oss-adoption", "Open-source dependency truth."],
            ["Cost and effort reduction", "/api/cost-effort-reduction", "How the app reduces decision/setup/rework burden."],
            ["Field utilities", "/api/field-utilities", "Field-level helpers and drilldowns."],
            ["OpenRouter model routing", "/api/ai/model-policy", "OpenRouter free/modest/expensive model policy, auto reasoning, and paid-tier approval boundary."],
            ["Source index", "/api/source-index/status", "NAS source index status and approved exclusions."],
        ],
        [1.75, 2.0, 2.75],
    )


def add_openrouter_story(doc: Document) -> None:
    heading(doc, "11. OpenRouter Reasoning Ladder Story", 1)
    paragraph(
        doc,
        "The user does not normally start on an AI provider screen. The OpenRouter path appears through Admin / Proof -> OpenRouter routing or through any future AI feature that needs model routing. The app first explains which tier it would use, then only performs a live provider call when explicit execution and approval gates are satisfied.",
    )
    table(
        doc,
        ["Tier", "Visible label", "Reasoning level", "Default model", "Approval story"],
        [
            ["free", "Free models first", "low / routine / cheap", "openrouter/free", "Preferred first path. A live call still requires OPENROUTER_API_KEY and execute=true, but it does not require paid-tier approval."],
            ["modest", "A little bit expensive", "medium / normal / moderate", "openai/gpt-5.1", "Used when the task needs stronger reasoning than the free tier. Live call requires execute=true and allow_paid_provider=true."],
            ["expensive", "Full expensive models", "high / deep / hard", "openai/gpt-5.5", "Reserved for deep audits, architecture, privacy/security, root cause work, product-bible synthesis, or high-value reasoning. Live call requires execute=true and allow_paid_provider=true."],
        ],
        [0.85, 1.35, 1.35, 1.35, 2.1],
    )
    flow(
        doc,
        "OpenRouter route-only flow",
        [
            "User clicks Admin / Proof -> OpenRouter routing",
            "App calls /api/ai/model-policy",
            "User or feature sends task text to /api/ai/route",
            "App infers free, modest, or expensive from reasoning level and task keywords",
            "UI shows selected model, paid status, and approval requirement",
            "No provider call occurs",
        ],
    )
    flow(
        doc,
        "OpenRouter live-call flow",
        [
            "User or future feature asks for AI completion",
            "App calls /api/ai/complete",
            "If execute is not true, app returns route recommendation only",
            "If selected tier is modest or expensive, allow_paid_provider must be true",
            "OPENROUTER_API_KEY stays server-side",
            "Only then does app call OpenRouter",
            "Response returns with provider_called true and selected tier proof",
        ],
    )
    callout(
        doc,
        "AI provider boundary",
        "OpenRouter is the chosen AI provider for this product path. The current app keeps local/internal features usable without any live provider call, and the UI presents model policy as proof before execution.",
        fill=GREEN,
    )


def add_runtime_appendix(doc: Document) -> None:
    heading(doc, "12. Runtime, Backend, And Data Story", 1)
    table(
        doc,
        ["Runtime item", "Value / route", "Story truth"],
        [
            ["Direct URL", "http://127.0.0.1:5420/", "Local app front door."],
            ["AMTL URL", "http://amtl/ck-life-os/", "Friendly gateway route."],
            ["Alias", "http://amtl/lifeos/", "Friendly alias route."],
            ["Health", "/health and /api/health", "Confirms runtime ok, port 5420, no gamification, runtime independence."],
            ["Practice receipts", "Local JSONL", "Local practice/system receipts, no external send."],
            ["Encrypted journal", "Local encrypted JSONL", "Private title/content encrypted at rest."],
            ["Inner Work", "Local encrypted JSONL", "Private guided sessions encrypted at rest."],
            ["Innovation catalogue", "Generated local catalogue", "500 field/navigation patterns; synthetic, not lived history."],
            ["RAG source drafts", "Local JSONL", "Internal/external source references plus optional pasted source text; no external fetch."],
            ["NAS source index", "Local metadata/text-snippet JSONL", "Internal source truth with exactly two approved exclusions."],
            ["OpenRouter", "/api/ai/model-policy; /api/ai/route; /api/ai/complete", "OpenRouter free first, modest paid, then full expensive. Live calls require key and execute=true; modest/full tiers also require allow_paid_provider=true."],
        ],
        [1.6, 2.45, 2.45],
    )
    flow(
        doc,
        "Generic click-to-proof flow",
        [
            "User clicks menu item",
            "App opens screen",
            "User clicks tab/button or enters field",
            "App calls backend route",
            "Local data/proof/result is generated",
            "Middle tab updates",
            "Right context rail updates",
            "Evidence remains local unless explicitly approved elsewhere",
        ],
    )
    heading(doc, "13. Dependency-Down And Boundary Story", 1)
    paragraph(
        doc,
        "If PostgreSQL, scheduler, OpenRouter, NAS, Workshop, Beast, or another AMTL product is down, the core Practice desk, Guide me, Encrypted journal, Inner work, Reports, and Evidence screens remain usable. Optional dependencies are labelled as optional, fallback, or unavailable; they are not hidden as fake success.",
    )
    table(
        doc,
        ["Dependency", "Core required?", "Fallback story"],
        [
            ["PostgreSQL / pgvector", "No", "Core uses local JSONL and generated/local data. RAG source search uses local drafts and NAS index, not pgvector-required."],
            ["Scheduler", "No", "Manual practice recording remains available."],
            ["OpenRouter", "No", "Local fallback explains routing. Provider_called remains false without key/execution/approval."],
            ["NAS source index", "No for core practice UI", "RAG / Sources screen can show not_built or access issues; practice flow remains usable."],
            ["Other AMTL apps", "No", "Handoffs are local staged gates only."],
        ],
        [1.7, 1.4, 3.4],
    )


def add_final_checklist(doc: Document) -> None:
    heading(doc, "14. Complete Screen And Button Checklist", 1)
    table(
        doc,
        ["Section", "Tabs / panels", "Buttons / fields covered"],
        [
            ["Top bar", "AMTL logo, CK Life OS logo/name, today's signal, date/time/status, AMTL seal", "No buttons; orientation only."],
            ["Practice desk", "Practice tabs; Today's truth", "Practice note; Make today gentler; Explain first step; Practices; Scores; Streaks; Lens items; Record; Explain prompt; Details."],
            ["Guide me", "Start; Practice 6W", "Practice; Ideas; Reports; Evidence."],
            ["Innovation lens", "Filter; Results; Detail", "Practice; Situation; Open matching innovations; Open {idea_id}."],
            ["Encrypted journal", "Write; Entries; Detail", "Practice; Title; Encrypted journal entry; Save encrypted entry; Refresh journal; Open local decrypt."],
            ["Inner work", "Start; Sessions; Detail", "Mode; Intention; Inner work text input; Voice transcript input; Start guided session; Load sessions; Open session."],
            ["RAG / Sources", "RAG status; Add source; Search; Results", "Explain RAG access; Refresh source status; Build local index; Source type; Title; Path or URL; RAG source notes; RAG source text; Approve one-time external URL fetch into local RAG; Stage source for approval; Search sources; Search."],
            ["Reports and gates", "Summary; Synthetic data", "Open local report; Export markdown summary; Stage Elaine local gate; Check boundary; Preview cleanup; Run cleanup verification."],
            ["Evidence summaries", "Checks; Output", "Product Bible matrix; R2D2 equivalent; Dependency fallback; Data truth; Button/drilldown truth; OSS adoption; Cost and effort reduction; Field utilities; OpenRouter model routing; Source index."],
            ["OpenRouter routing", "Admin / Proof direct button; Evidence summaries proof button", "Free models first; A little bit expensive; Full expensive models; auto reasoning; execute=true; allow_paid_provider=true for paid tiers; server-side OPENROUTER_API_KEY boundary."],
            ["Right context rail", "Next action; Selected detail; 6W / ELI10; Receipt / Evidence", "Collapsed disclosure panels; updates when the user clicks."],
        ],
        [1.4, 2.35, 2.75],
    )
    paragraph(doc, "Plain-English ending", bold=True, color=BLUE, before=4, after=2)
    paragraph(
        doc,
        "CK / Life OS is a calm local/internal practice operating system. The user can practice, reflect privately, do Inner Work, use innovation helpers, inspect sources, create local reports, and review proof without the app becoming noisy, performative, or externally mutating anything.",
        after=0,
    )


def build() -> None:
    doc = setup_doc()
    add_cover(doc)
    add_global_orientation(doc)
    add_practice_story(doc)
    add_guide_and_lens(doc)
    add_private_story(doc)
    add_sources_reports_evidence(doc)
    add_openrouter_story(doc)
    add_runtime_appendix(doc)
    add_final_checklist(doc)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.text = f"CK / Life OS Complete Story-Flow Guide {DOC_VERSION}"
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            set_run(run, size=8.5, color="6B7785")
        footer = section.footer.paragraphs[0]
        footer.text = f"Local/internal product story map - generated for Mani - {DOC_VERSION}"
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in footer.runs:
            set_run(run, size=8.5, color="6B7785")

    doc.core_properties.title = f"CK / Life OS Complete Story-Flow Guide {DOC_VERSION}"
    doc.core_properties.subject = "Complete story-like flow map for all CK / Life OS modules, screens, panels, buttons, and runtime proof"
    doc.core_properties.author = "AMTL Codex"
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
